"""
Enhanced rule-based resume parser - MVP-2
Extracts all resume sections with bullet points and achievements
Structured for easy AI integration in Phase 2
"""
import re
from typing import Optional, List, Tuple
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Common skill keywords to look for
COMMON_SKILLS = [
    "python", "javascript", "typescript", "react", "vue", "angular", "node.js", "nodejs",
    "java", "c++", "c#", "go", "golang", "rust", "ruby", "php", "swift", "kotlin",
    "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "jenkins", "ci/cd",
    "git", "linux", "bash", "rest api", "graphql", "microservices",
    "machine learning", "ml", "ai", "deep learning", "tensorflow", "pytorch",
    "data analysis", "pandas", "numpy", "data science", "statistics",
    "html", "css", "sass", "tailwind", "bootstrap",
    "agile", "scrum", "jira", "product management", "project management",
    "figma", "sketch", "ui/ux", "design", "photoshop",
    "communication", "leadership", "problem solving", "teamwork",
    "excel", "powerpoint", "word", "google analytics",
    "fastapi", "django", "flask", "express", "spring boot",
    "nextjs", "next.js", "gatsby", "webpack", "vite"
]

# Regex patterns
EMAIL_PATTERN = re.compile(r'[\w\.-]+@[\w\.-]+\.\w+')
PHONE_PATTERN = re.compile(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
LINKEDIN_PATTERN = re.compile(r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9_-]+)', re.IGNORECASE)
GITHUB_PATTERN = re.compile(r'(?:github\.com/|github:?\s*)([a-zA-Z0-9_-]+)', re.IGNORECASE)
URL_PATTERN = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+')
DATE_PATTERN = re.compile(r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[.\s]+\d{4}|20\d{2}', re.IGNORECASE)
DATE_RANGE_PATTERN = re.compile(r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[.\s]*\d{4}|20\d{2})\s*[-–to]+\s*((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[.\s]*\d{4}|20\d{2}|[Pp]resent|[Cc]urrent|[Nn]ow)', re.IGNORECASE)

# Section headers
SECTION_HEADERS = {
    'experience': ['experience', 'work history', 'employment', 'professional experience', 'work experience', 'career history', 'professional background'],
    'education': ['education', 'academic', 'academics', 'qualifications', 'educational background'],
    'skills': ['skills', 'technical skills', 'competencies', 'technologies', 'tools', 'expertise'],
    'projects': ['projects', 'personal projects', 'side projects', 'portfolio'],
    'certifications': ['certifications', 'certificates', 'professional development', 'licenses'],
    'summary': ['summary', 'objective', 'profile', 'about me', 'professional summary', 'career objective'],
    'languages': ['languages', 'language skills']
}

# Bullet point markers
BULLET_MARKERS = ['•', '●', '○', '◦', '▪', '▸', '►', '-', '*', '→', '»', '›']


class ResumeParser:
    """
    Enhanced rule-based resume parser with comprehensive section extraction
    """
    
    def __init__(self, ai_client=None):
        self.ai_client = ai_client
    
    def parse_file(self, filepath: str) -> dict:
        """Parse any supported file type (PDF, DOCX, TXT) and extract information"""
        filepath_lower = filepath.lower()
        
        if filepath_lower.endswith('.pdf'):
            return self.parse_pdf(filepath)
        elif filepath_lower.endswith('.docx') or filepath_lower.endswith('.doc'):
            return self.parse_docx(filepath)
        elif filepath_lower.endswith('.txt'):
            return self.parse_txt(filepath)
        else:
            # Try as text file
            return self.parse_txt(filepath)
    
    def parse_pdf(self, filepath: str) -> dict:
        """Parse PDF file and extract information"""
        try:
            text = self._extract_text_from_pdf(filepath)
            if not text:
                return {"error": "Could not extract text from PDF"}
            # Clean the extracted text
            text = self._clean_extracted_text(text)
            return self.parse_text(text)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return {"error": str(e)}
    
    def parse_docx(self, filepath: str) -> dict:
        """Parse DOCX file and extract information"""
        try:
            text = self._extract_text_from_docx(filepath)
            if not text:
                return {"error": "Could not extract text from DOCX"}
            text = self._clean_extracted_text(text)
            return self.parse_text(text)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return {"error": str(e)}
    
    def parse_txt(self, filepath: str) -> dict:
        """Parse TXT file and extract information"""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            if not text:
                return {"error": "Could not read text file"}
            return self.parse_text(text)
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            return {"error": str(e)}
    
    def _extract_text_from_docx(self, filepath: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return None
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean up PDF extracted text - fix common extraction issues"""
        # Add spaces between camelCase words that got merged
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Add spaces after periods/commas if missing
        text = re.sub(r'([.,])([A-Za-z])', r'\1 \2', text)
        
        # Fix common concatenated patterns
        text = re.sub(r'([a-z])([A-Z][a-z])', r'\1 \2', text)
        
        # Clean up multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Ensure newlines are preserved
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def _extract_text_from_pdf(self, filepath: str) -> Optional[str]:
        """Extract text from PDF using pdfplumber or fallback to PyMuPDF"""
        text = ""
        
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    # Use layout extraction for better spacing
                    page_text = page.extract_text(layout=True)
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except ImportError:
            logger.warning("pdfplumber not installed, trying PyMuPDF")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        try:
            import fitz
            doc = fitz.open(filepath)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            if text.strip():
                return text
        except ImportError:
            logger.warning("PyMuPDF not installed")
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {e}")
        
        return text if text.strip() else None
    
    def parse_text(self, text: str) -> dict:
        """Parse resume text and extract all sections"""
        result = {
            "name": None,
            "email": None,
            "phone": None,
            "location": None,
            "linkedin": None,
            "github": None,
            "portfolio": None,
            "summary": None,
            "objective": None,
            "skills": [],
            "education": [],
            "experience": [],
            "projects": [],
            "certifications": [],
            "languages": [],
            "interests": [],
            "total_years_experience": None
        }
        
        lines = [line.strip() for line in text.split('\n')]
        text_lower = text.lower()
        
        # Extract contact information
        result["email"] = self._extract_email(text)
        result["phone"] = self._extract_phone(text)
        result["name"] = self._extract_name(lines)
        result["linkedin"] = self._extract_linkedin(text)
        result["github"] = self._extract_github(text)
        result["portfolio"] = self._extract_portfolio(text)
        result["location"] = self._extract_location(lines)
        
        # Identify section boundaries
        sections = self._identify_sections(lines)
        
        # Extract each section
        result["summary"] = self._extract_summary_section(lines, sections)
        result["skills"] = self._extract_skills(text_lower, lines, sections)
        result["education"] = self._extract_education(lines, sections)
        result["experience"] = self._extract_experience(lines, sections)
        result["projects"] = self._extract_projects(lines, sections)
        result["certifications"] = self._extract_certifications(lines, sections)
        result["languages"] = self._extract_languages(lines, sections)
        
        # Calculate total years of experience
        result["total_years_experience"] = self._calculate_experience_years(result["experience"])
        
        return result
    
    def _extract_email(self, text: str) -> Optional[str]:
        match = EMAIL_PATTERN.search(text)
        return match.group() if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        match = PHONE_PATTERN.search(text)
        return match.group() if match else None
    
    def _extract_name(self, lines: List[str]) -> Optional[str]:
        """Extract name from first few lines"""
        for line in lines[:5]:
            line = line.strip()
            if line and len(line) > 2 and len(line) < 50:
                # Skip lines with common non-name patterns
                if not EMAIL_PATTERN.search(line) and not PHONE_PATTERN.search(line):
                    if not any(kw in line.lower() for kw in ['resume', 'cv', 'curriculum', 'http', 'www', '@']):
                        # Should contain mostly letters and spaces
                        if sum(c.isalpha() or c.isspace() for c in line) / len(line) > 0.8:
                            return line
        return None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        match = LINKEDIN_PATTERN.search(text)
        if match:
            return f"linkedin.com/in/{match.group(1)}"
        return None
    
    def _extract_github(self, text: str) -> Optional[str]:
        match = GITHUB_PATTERN.search(text)
        if match:
            return f"github.com/{match.group(1)}"
        return None
    
    def _extract_portfolio(self, text: str) -> Optional[str]:
        urls = URL_PATTERN.findall(text)
        for url in urls:
            url_lower = url.lower()
            if 'linkedin' not in url_lower and 'github' not in url_lower:
                if any(kw in url_lower for kw in ['portfolio', 'personal', '.me', '.io', 'behance', 'dribbble']):
                    return url
        return None
    
    def _extract_location(self, lines: List[str]) -> Optional[str]:
        """Extract location from header area"""
        location_patterns = [
            re.compile(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2})', re.MULTILINE),  # City, ST
            re.compile(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z][a-z]+)', re.MULTILINE),  # City, State
        ]
        for line in lines[:10]:
            for pattern in location_patterns:
                match = pattern.search(line)
                if match:
                    return match.group(1)
        return None
    
    def _identify_sections(self, lines: List[str]) -> dict:
        """Identify section boundaries in the resume"""
        sections = {}
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            for section_type, keywords in SECTION_HEADERS.items():
                for kw in keywords:
                    if line_lower == kw or line_lower.startswith(kw + ':') or line_lower.startswith(kw + ' '):
                        if section_type not in sections:
                            sections[section_type] = i
                        break
        return sections
    
    def _get_section_lines(self, lines: List[str], sections: dict, section_name: str, default_limit: int = 20) -> List[str]:
        """Get lines belonging to a specific section"""
        if section_name not in sections:
            return []
        
        start = sections[section_name] + 1
        # Find where the next section starts
        next_section_start = len(lines)
        for name, idx in sections.items():
            if idx > sections[section_name] and idx < next_section_start:
                next_section_start = idx
        
        return lines[start:min(next_section_start, start + default_limit)]
    
    def _extract_summary_section(self, lines: List[str], sections: dict) -> Optional[str]:
        """Extract summary/objective section"""
        section_lines = self._get_section_lines(lines, sections, 'summary', 10)
        if section_lines:
            summary_text = ' '.join([l for l in section_lines if l and len(l) > 10])
            return summary_text[:500] if summary_text else None
        return None
    
    def _extract_skills(self, text_lower: str, lines: List[str], sections: dict) -> List[str]:
        """Extract skills from text and skills section"""
        found_skills = set()
        
        # Extract from common skills list
        for skill in COMMON_SKILLS:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill.title())
        
        # Also extract from skills section if present
        section_lines = self._get_section_lines(lines, sections, 'skills', 15)
        for line in section_lines:
            # Split by common delimiters
            items = re.split(r'[,;|•●○▪►\-\*]', line)
            for item in items:
                item = item.strip()
                if item and 2 < len(item) < 30:
                    found_skills.add(item)
        
        return list(found_skills)[:30]
    
    def _extract_education(self, lines: List[str], sections: dict) -> List[dict]:
        """Extract education entries with achievements"""
        education = []
        section_lines = self._get_section_lines(lines, sections, 'education', 25)
        
        current_entry = None
        for line in section_lines:
            if not line:
                continue
            
            # Check if this is a new entry (usually has institution name or degree)
            is_new_entry = False
            line_lower = line.lower()
            
            if any(kw in line_lower for kw in ['university', 'college', 'institute', 'school', 'bachelor', 'master', 'phd', 'mba', 'b.s.', 'b.a.', 'm.s.', 'm.a.']):
                is_new_entry = True
            
            if is_new_entry:
                if current_entry:
                    education.append(current_entry)
                
                # Parse date range
                start_date, end_date = self._parse_date_range(line)
                
                current_entry = {
                    "institution": self._clean_line(line),
                    "degree": self._extract_degree(line),
                    "field": self._extract_field(line),
                    "start_date": start_date,
                    "end_date": end_date,
                    "year": self._extract_year(line),
                    "gpa": self._extract_gpa(line),
                    "achievements": []
                }
            elif current_entry:
                # Check if it's a bullet point
                if self._is_bullet_point(line):
                    current_entry["achievements"].append(self._clean_bullet(line))
                elif 'gpa' in line_lower:
                    current_entry["gpa"] = self._extract_gpa(line)
                elif not current_entry.get("degree"):
                    current_entry["degree"] = self._extract_degree(line)
        
        if current_entry:
            education.append(current_entry)
        
        return education[:5]
    
    def _extract_experience(self, lines: List[str], sections: dict) -> List[dict]:
        """
        Extract work experience with proper structure:
        - Line with date = Job Title line
        - Next non-bullet line = Company/Location line  
        - Bullet points = Achievements
        """
        experience = []
        section_lines = self._get_section_lines(lines, sections, 'experience', 100)
        
        if not section_lines:
            return experience
        
        current_entry = None
        expecting_company = False
        
        for i, line in enumerate(section_lines):
            if not line:
                continue
            
            is_bullet = self._is_bullet_point(line)
            has_date = bool(DATE_RANGE_PATTERN.search(line) or DATE_PATTERN.search(line))
            
            # Line with date = new job entry (this is the JOB TITLE line)
            if has_date and not is_bullet:
                # Save previous entry
                if current_entry:
                    experience.append(current_entry)
                
                start_date, end_date = self._parse_date_range(line)
                is_current = end_date and end_date.lower() in ['present', 'current', 'now'] if end_date else False
                
                # Remove the date from the line to get the job title
                title_text = DATE_RANGE_PATTERN.sub('', line)
                title_text = DATE_PATTERN.sub('', title_text)
                title_text = title_text.strip().strip('|–-,').strip()
                
                current_entry = {
                    "company": None,  # Will be filled from next line
                    "title": title_text if title_text else None,
                    "location": None,
                    "start_date": start_date,
                    "end_date": end_date if not is_current else "Present",
                    "duration": f"{start_date} - {end_date}" if start_date and end_date else None,
                    "is_current": is_current,
                    "description": None,
                    "achievements": []
                }
                expecting_company = True
                
            elif current_entry and expecting_company and not is_bullet:
                # This line should be the company name (and possibly location)
                company_line = line.strip()
                
                # Extract location if present (usually "Company, City, State" or "Company City, State")
                location = self._extract_job_location(company_line)
                if location:
                    # Remove location from company name
                    company_name = company_line.replace(location, '').strip().strip(',').strip()
                else:
                    company_name = company_line
                
                current_entry["company"] = company_name
                current_entry["location"] = location
                expecting_company = False
                
            elif current_entry and is_bullet:
                # Bullet point = achievement
                achievement = self._clean_bullet(line)
                if achievement and len(achievement) > 10:
                    current_entry["achievements"].append(achievement)
                expecting_company = False
                
            elif current_entry and not is_bullet and len(line) > 20:
                # Non-bullet continuation line - append to last achievement or description
                if current_entry["achievements"]:
                    # Append to last achievement (likely a wrapped line)
                    current_entry["achievements"][-1] += " " + line
                elif not current_entry.get("company"):
                    # Might be the company line
                    current_entry["company"] = line
                    expecting_company = False
        
        if current_entry:
            experience.append(current_entry)
        
        # Post-process: If company is None, swap with title
        for exp in experience:
            if not exp.get("company") and exp.get("title"):
                # Title might actually be the company
                exp["company"] = exp["title"]
                exp["title"] = None
        
        # Sort by start date (most recent first)
        experience = self._sort_by_date(experience)
        
        return experience[:10]
    
    def _extract_projects(self, lines: List[str], sections: dict) -> List[dict]:
        """Extract projects section"""
        projects = []
        section_lines = self._get_section_lines(lines, sections, 'projects', 30)
        
        current_project = None
        
        for line in section_lines:
            if not line:
                continue
            
            # New project usually has a name (shorter line, possibly with link or date)
            is_new_project = False
            if not self._is_bullet_point(line) and len(line) < 80:
                if URL_PATTERN.search(line) or DATE_PATTERN.search(line) or line[0].isupper():
                    is_new_project = True
            
            if is_new_project:
                if current_project:
                    projects.append(current_project)
                
                urls = URL_PATTERN.findall(line)
                
                current_project = {
                    "name": self._clean_line(line),
                    "description": None,
                    "technologies": [],
                    "url": urls[0] if urls else None,
                    "achievements": []
                }
            elif current_project:
                if self._is_bullet_point(line):
                    bullet_text = self._clean_bullet(line)
                    if bullet_text:
                        current_project["achievements"].append(bullet_text)
                        # Extract technologies from bullet points
                        techs = self._extract_technologies(bullet_text)
                        current_project["technologies"].extend(techs)
                elif not current_project.get("description"):
                    current_project["description"] = line[:200]
        
        if current_project:
            projects.append(current_project)
        
        # Remove duplicate technologies
        for project in projects:
            project["technologies"] = list(set(project["technologies"]))[:10]
        
        return projects[:8]
    
    def _extract_certifications(self, lines: List[str], sections: dict) -> List[dict]:
        """Extract certifications"""
        certifications = []
        section_lines = self._get_section_lines(lines, sections, 'certifications', 15)
        
        for line in section_lines:
            if not line or len(line) < 5:
                continue
            
            cert = {
                "name": self._clean_line(line),
                "issuer": self._extract_issuer(line),
                "date": self._extract_year(line),
                "credential_id": None,
                "url": None
            }
            
            urls = URL_PATTERN.findall(line)
            if urls:
                cert["url"] = urls[0]
            
            certifications.append(cert)
        
        return certifications[:10]
    
    def _extract_languages(self, lines: List[str], sections: dict) -> List[str]:
        """Extract languages"""
        languages = []
        section_lines = self._get_section_lines(lines, sections, 'languages', 10)
        
        for line in section_lines:
            # Split by common delimiters
            items = re.split(r'[,;|•●]', line)
            for item in items:
                item = item.strip()
                # Remove proficiency indicators
                item = re.sub(r'\s*\([^)]+\)\s*', '', item)
                item = re.sub(r'\s*-\s*(Native|Fluent|Advanced|Intermediate|Basic).*$', '', item, flags=re.IGNORECASE)
                if item and 2 < len(item) < 30:
                    languages.append(item)
        
        return languages[:10]
    
    def _is_bullet_point(self, line: str) -> bool:
        """Check if line starts with a bullet marker"""
        line = line.strip()
        if not line:
            return False
        return line[0] in BULLET_MARKERS or line.startswith('- ') or line.startswith('* ')
    
    def _clean_bullet(self, line: str) -> str:
        """Remove bullet markers and clean the text"""
        line = line.strip()
        for marker in BULLET_MARKERS:
            if line.startswith(marker):
                line = line[len(marker):].strip()
                break
        return line
    
    def _clean_line(self, line: str) -> str:
        """Clean line by removing dates and extra whitespace"""
        line = DATE_RANGE_PATTERN.sub('', line)
        line = DATE_PATTERN.sub('', line)
        line = re.sub(r'\s+', ' ', line)
        return line.strip()[:150]
    
    def _parse_date_range(self, text: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract start and end dates from text"""
        match = DATE_RANGE_PATTERN.search(text)
        if match:
            return match.group(1), match.group(2)
        
        # Single date
        match = DATE_PATTERN.search(text)
        if match:
            return match.group(), None
        
        return None, None
    
    def _extract_year(self, text: str) -> Optional[str]:
        """Extract year or date range from text"""
        match = DATE_RANGE_PATTERN.search(text)
        if match:
            return f"{match.group(1)} - {match.group(2)}"
        
        year_pattern = re.compile(r'20\d{2}')
        match = year_pattern.search(text)
        return match.group() if match else None
    
    def _extract_degree(self, text: str) -> Optional[str]:
        """Extract degree from text"""
        degree_patterns = [
            r"(Bachelor(?:'s)?\s+(?:of\s+)?(?:Science|Arts|Engineering|Business))",
            r"(Master(?:'s)?\s+(?:of\s+)?(?:Science|Arts|Business|Engineering))",
            r"(Ph\.?D\.?|Doctor(?:ate)?)",
            r"(MBA|M\.B\.A\.)",
            r"(B\.S\.|B\.A\.|M\.S\.|M\.A\.)",
            r"(Associate(?:'s)?\s+(?:of\s+)?(?:Science|Arts))"
        ]
        for pattern in degree_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        return None
    
    def _extract_field(self, text: str) -> Optional[str]:
        """Extract field of study"""
        field_patterns = [
            r"in\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
            r"(?:Science|Arts|Engineering)\s+in\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)"
        ]
        for pattern in field_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        return None
    
    def _extract_gpa(self, text: str) -> Optional[str]:
        """Extract GPA from text"""
        gpa_pattern = re.compile(r'(?:GPA|Grade)[\s:]*(\d+\.?\d*(?:/\d+\.?\d*)?)', re.IGNORECASE)
        match = gpa_pattern.search(text)
        return match.group(1) if match else None
    
    def _split_company_title(self, text: str) -> Tuple[str, Optional[str]]:
        """Split company and title from a line"""
        # Remove dates first
        text = DATE_RANGE_PATTERN.sub('', text)
        text = DATE_PATTERN.sub('', text)
        
        # Common job titles to look for
        job_titles = [
            r'Software Engineer(?:ing)?\s*(?:Intern|Specialist|Lead|Manager|Sr\.?|Jr\.?|Senior|Junior)?',
            r'(?:Sr\.?|Jr\.?|Senior|Junior)?\s*Software Engineer(?:ing)?(?:\s*Intern)?',
            r'(?:Sr\.?|Jr\.?|Senior|Junior)?\s*Developer',
            r'(?:Sr\.?|Jr\.?|Senior|Junior)?\s*(?:Full[\s-]?Stack|Frontend|Backend|Full Stack)\s*(?:Developer|Engineer)',
            r'Product\s*(?:Manager|Designer|Owner)',
            r'Data\s*(?:Scientist|Engineer|Analyst)',
            r'(?:Machine Learning|ML|AI)\s*Engineer',
            r'DevOps\s*Engineer',
            r'(?:Engineering|Technical|Tech)\s*(?:Manager|Lead)',
            r'(?:UX|UI|UX/UI)\s*(?:Designer|Engineer)',
            r'(?:QA|Quality)\s*(?:Engineer|Analyst|Tester)',
            r'(?:Project|Program)\s*Manager',
            r'(?:Solutions?|Cloud)\s*(?:Architect|Engineer)',
            r'(?:Mobile|iOS|Android)\s*(?:Developer|Engineer)',
            r'(?:Web|Frontend|Backend)\s*Developer',
            r'Intern(?:ship)?',
        ]
        
        # Try to find job title pattern
        for title_pattern in job_titles:
            match = re.search(title_pattern, text, re.IGNORECASE)
            if match:
                title = match.group(0).strip()
                # The company is everything before or after the title
                company = text.replace(match.group(0), '').strip()
                company = re.sub(r'^[\s,\-–|]+|[\s,\-–|]+$', '', company)
                if company and title:
                    return company[:100], title[:100]
        
        # Common separators
        separators = ['|', ' - ', ' – ', ' at ', ' @ ']
        for sep in separators:
            if sep in text:
                parts = text.split(sep, 1)
                return parts[0].strip()[:100], parts[1].strip()[:100] if len(parts) > 1 else None
        
        return text.strip()[:100], None
    
    def _extract_job_location(self, text: str) -> Optional[str]:
        """Extract location from job line"""
        location_pattern = re.compile(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2})')
        match = location_pattern.search(text)
        return match.group(1) if match else None
    
    def _extract_issuer(self, text: str) -> Optional[str]:
        """Extract certification issuer"""
        issuer_patterns = [
            r'(?:from|by|issued by)\s+([A-Z][a-zA-Z\s]+)',
            r'-\s*([A-Z][a-zA-Z\s]+?)(?:\s*(?:20\d{2}|$))'
        ]
        for pattern in issuer_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()[:50]
        return None
    
    def _extract_technologies(self, text: str) -> List[str]:
        """Extract technology names from text"""
        text_lower = text.lower()
        techs = []
        for skill in COMMON_SKILLS:
            if skill in text_lower:
                techs.append(skill.title())
        return techs
    
    def _sort_by_date(self, entries: List[dict]) -> List[dict]:
        """Sort entries by start date (most recent first)"""
        def date_key(entry):
            start = entry.get("start_date", "") or ""
            # Extract year for sorting
            year_match = re.search(r'20\d{2}', start)
            return int(year_match.group()) if year_match else 0
        
        return sorted(entries, key=date_key, reverse=True)
    
    def _calculate_experience_years(self, experience: List[dict]) -> Optional[float]:
        """Calculate total years of experience"""
        import datetime as dt
        total_months = 0
        current_year = dt.datetime.now().year
        
        for exp in experience:
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            
            start_year_match = re.search(r'20\d{2}', start) if start else None
            end_year_match = re.search(r'20\d{2}', end) if end else None
            
            start_year = int(start_year_match.group()) if start_year_match else None
            
            if end and end.lower() in ['present', 'current', 'now']:
                end_year = current_year
            else:
                end_year = int(end_year_match.group()) if end_year_match else start_year
            
            if start_year and end_year:
                total_months += (end_year - start_year + 1) * 12
        
        return round(total_months / 12, 1) if total_months > 0 else None


# Singleton instance
parser = ResumeParser()


def parse_resume(filepath: str) -> dict:
    """Convenience function for parsing resumes - supports PDF, DOCX, TXT"""
    return parser.parse_file(filepath)
